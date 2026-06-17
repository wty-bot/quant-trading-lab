from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT = ROOT / "reports" / "final_report.md"
OUT = ROOT / "final_submission_assets" / "report" / "A股多因子量化交易系统课程报告.docx"
MD_OUT = ROOT / "final_submission_assets" / "report" / "A股多因子量化交易系统课程报告.md"


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12, bold: bool = False) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, alignment=None, size: float = 12) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    if alignment is not None:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_paragraph(document: Document, text: str, size: float = 12, alignment=None, bold: bool = False) -> None:
    para = document.add_paragraph()
    para.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_heading(document: Document, text: str, level: int) -> None:
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    if level == 0:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(run, east_asia="黑体", size=16, bold=True)
    elif level == 1:
        set_run_font(run, east_asia="黑体", size=15, bold=True)
    else:
        set_run_font(run, east_asia="黑体", size=14, bold=True)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def add_markdown_table(document: Document, lines: list[str], caption: str | None) -> None:
    rows = []
    for line in lines:
        parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(part) <= {"-", ":"} for part in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, cell in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], cell, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row[: len(cells)]):
            set_cell_text(cells[i], cell_text)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "999999"},
                bottom={"val": "single", "sz": "4", "color": "999999"},
                left={"val": "single", "sz": "4", "color": "CCCCCC"},
                right={"val": "single", "sz": "4", "color": "CCCCCC"},
            )
    if caption:
        add_caption(document, caption)


def add_caption(document: Document, caption: str) -> None:
    add_paragraph(document, caption, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def is_caption(line: str, prefix: str) -> bool:
    return re.match(rf"^{re.escape(prefix)}\s*\d+\s{{2,}}\S+", line) is not None


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    MD_OUT.write_text(REPORT.read_text(encoding="utf-8"), encoding="utf-8")

    document = Document()
    document.core_properties.title = "A股多因子量化交易系统课程报告"
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.05)
    section.right_margin = Inches(1.05)

    lines = REPORT.read_text(encoding="utf-8").splitlines()
    table_buffer: list[str] = []
    waiting_for_table_caption = False
    waiting_for_figure_caption = False
    in_code_block = False

    def flush_table() -> None:
        nonlocal table_buffer, waiting_for_table_caption
        if table_buffer:
            add_markdown_table(document, table_buffer, None)
            table_buffer = []
            waiting_for_table_caption = True

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code_block = not in_code_block
            continue
        if in_code_block:
            add_paragraph(document, line, size=10.5)
            continue
        if line.startswith("|") and line.endswith("|"):
            waiting_for_table_caption = False
            table_buffer.append(line)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            add_heading(document, line[2:].strip(), 0)
        elif line.startswith("## "):
            add_heading(document, line[3:].strip(), 1)
        elif line.startswith("### "):
            add_heading(document, line[4:].strip(), 2)
        elif waiting_for_table_caption and is_caption(line, "表"):
            add_caption(document, line)
            waiting_for_table_caption = False
        elif is_caption(line, "图"):
            if waiting_for_figure_caption:
                add_caption(document, line)
                waiting_for_figure_caption = False
            else:
                add_paragraph(document, line)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            waiting_for_table_caption = False
            rel = line[line.index("(") + 1 : -1]
            image_path = (REPORT.parent / rel).resolve()
            if image_path.exists():
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(str(image_path), width=Inches(5.9))
                waiting_for_figure_caption = True
        else:
            waiting_for_table_caption = False
            waiting_for_figure_caption = False
            add_paragraph(document, line)

    flush_table()
    document.save(OUT)


if __name__ == "__main__":
    main()
